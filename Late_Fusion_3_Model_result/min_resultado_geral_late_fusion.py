import numpy as np
import scipy.io as sio
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sklearn.metrics import (
    confusion_matrix, f1_score, matthews_corrcoef,
    cohen_kappa_score, roc_curve, auc,
    precision_recall_curve, average_precision_score,
)

# =========================================================
# CONFIGURAÇÃO DAS COMBINAÇÕES DE LATE FUSION (MINIMUM POST-HOC)
# Os arquivos .mat foram gerados pelo minimum_late_fusion_post_hoc.py
# e salvos em late_fusion_post_hoc/fusion_3_model/minimum, com os seguintes padrões
# de nome:
#   predictions_test _ {modelo1}_{modelo2}.mat   (obs.: " _ " com espaços)
#   test_predict_{modelo1}_{modelo2}.mat
#   test_labels.mat                              (único, comum a todas)
# =========================================================
save_path_late_fusion = "/media/gledson/Dados1/Classificação_GastroIntestinal/classificar_original/late_fusion_post_hoc/fusion_3_model/minimum"

nomes_modelos = {
    'convnextv2':      'ConvNeXtV2-L',
    'dinov2':          'DINOv2',
    'efficientnetv2m': 'EfficientNetV2-M',
}

m1 = 'convnextv2'
m2 = 'dinov2'
m3 = 'efficientnetv2m'

modelos = [{
        'nome': f'Late Fusion (Minimum) — {nomes_modelos[m1]} + {nomes_modelos[m2]} + {nomes_modelos[m3]}',
        'path_predictions': save_path_late_fusion + '/predictions_test' + "_" + m1 + "_" + m2 + "_" + m3 + '.mat',
        'path_predict':     save_path_late_fusion + '/test_predict' + "_" + m1 + "_" + m2 + "_" + m3 + '.mat',
        'path_labels':      save_path_late_fusion + '/test_labels.mat',
    }]

save_path = '/media/gledson/Dados1/Classificação_GastroIntestinal/classificar_original/late_fusion_post_hoc/fusion_3_model/minimum/comparacao_minimum_late_fusion/'
os.makedirs(save_path, exist_ok=True)

class_names = ['lesion', 'normal']
colors_roc  = ['steelblue', 'crimson']
colors_pr   = ['steelblue', 'crimson']

# =========================================================
# FUNÇÕES AUXILIARES
# =========================================================
def calcular_metricas_por_classe(y_true, y_pred, class_idx):
    y_true_bin = (y_true == class_idx).astype(int)
    y_pred_bin = (y_pred == class_idx).astype(int)
    cm_bin = confusion_matrix(y_true_bin, y_pred_bin)
    TN, FP, FN, TP = cm_bin[0,0], cm_bin[0,1], cm_bin[1,0], cm_bin[1,1]
    total = TP + TN + FP + FN
    return {
        'TP':                       int(TP),
        'FP':                       int(FP),
        'FN':                       int(FN),
        'TN':                       int(TN),
        'Accuracy':                 (TP+TN)/total if total>0 else 0,
        'Sensitivity':              TP/(TP+FN) if (TP+FN)>0 else 0,
        'Specificity':              TN/(TN+FP) if (TN+FP)>0 else 0,
        'Precision':                TP/(TP+FP) if (TP+FP)>0 else 0,
        'FalsePositiveRate':        FP/(FP+TN) if (FP+TN)>0 else 0,
        'F1_score':                 f1_score(y_true_bin, y_pred_bin, zero_division=0),
        'MatthewsCorrelationCoef':  matthews_corrcoef(y_true_bin, y_pred_bin),
        'Kappa':                    cohen_kappa_score(y_true_bin, y_pred_bin),
    }


def threshold_otimo_pr(precision, recall, thresholds):
    f1s     = 2*(precision[:-1]*recall[:-1])/(precision[:-1]+recall[:-1]+1e-8)
    idx     = np.argmax(f1s)
    return thresholds[idx], precision[idx], recall[idx], f1s[idx], idx


def avaliar_modelo(cfg):
    """Carrega os .mat da combinação de late fusion, calcula todas as
    métricas e retorna dicionário completo."""
    tl = sio.loadmat(cfg['path_labels'])['test_labels'].ravel().astype(int)
    tp = sio.loadmat(cfg['path_predict'])['test_predict'].ravel().astype(int)
    pt = sio.loadmat(cfg['path_predictions'])['predictions_test']

    prob = {0: pt[:, 0], 1: pt[:, 1]}

    resultado = {'nome': cfg['nome'], 'metricas': {}, 'roc': {}, 'pr': {}}

    for i, cn in enumerate(class_names):
        resultado['metricas'][cn] = calcular_metricas_por_classe(tl, tp, i)

        yb = (tl == i).astype(int)

        # ROC
        fpr, tpr, thr_r = roc_curve(yb, prob[i])
        roc_auc          = auc(fpr, tpr)
        yi               = np.argmax(tpr - fpr)
        resultado['roc'][cn] = {
            'fpr': fpr, 'tpr': tpr, 'auc': roc_auc,
            'thr_opt': thr_r[yi], 'sens_opt': tpr[yi],
            'espec_opt': 1-fpr[yi], 'youden_idx': yi,
            'thr_arr': thr_r,
        }

        # PR
        prec, rec, thr_p = precision_recall_curve(yb, prob[i])
        ap                = average_precision_score(yb, prob[i])
        thr_opt, p_opt, r_opt, f1_opt, idx_opt = threshold_otimo_pr(prec, rec, thr_p)
        resultado['pr'][cn] = {
            'prec': prec, 'rec': rec, 'ap': ap,
            'thr_opt': thr_opt, 'prec_opt': p_opt,
            'rec_opt': r_opt, 'f1_opt': f1_opt,
            'idx_opt': idx_opt, 'baseline': yb.mean(),
        }

    return resultado


# =========================================================
# AVALIAR TODAS AS COMBINAÇÕES DE LATE FUSION
# =========================================================
print('Avaliando combinações de late fusion (minimum post-hoc)...')
resultados = [avaliar_modelo(cfg) for cfg in modelos]
print(f'{len(resultados)} combinação(ões) avaliada(s).\n')

# =========================================================
# FIGURA 1 — CURVAS ROC (todas as combinações, por classe)
# =========================================================
fig, axes = plt.subplots(1, 2, figsize=(14, 6))
line_styles = ['-', '--', '-.', ':']

for ci, cn in enumerate(class_names):
    ax = axes[ci]
    for mi, res in enumerate(resultados):
        r    = res['roc'][cn]
        ls   = line_styles[mi % len(line_styles)]
        ax.plot(r['fpr'], r['tpr'], lw=2, linestyle=ls,
                label=f'{res["nome"]} (AUC={r["auc"]:.4f})')
        ax.scatter(r['fpr'][r['youden_idx']], r['tpr'][r['youden_idx']],
                   zorder=5, s=60,
                   label=f'  Thr={r["thr_opt"]:.3f} Sens={r["sens_opt"]:.3f} Espec={r["espec_opt"]:.3f}')

    ax.plot([0,1],[0,1], color='gray', lw=1, linestyle='--', label='Aleatório')
    ax.set_xlabel('1 - Especificidade (FPR)', fontsize=11)
    ax.set_ylabel('Sensibilidade (TPR)',       fontsize=11)
    ax.set_title(f'Curva ROC — Classe "{cn}"', fontsize=12)
    ax.legend(loc='lower right', fontsize=7)
    ax.grid(True, alpha=0.3)
    ax.set_xlim([0,1]); ax.set_ylim([0,1.05])

plt.tight_layout()
plt.savefig(os.path.join(save_path, 'roc_comparacao_late_fusion.png'), dpi=600, bbox_inches='tight')
plt.show()

# =========================================================
# FIGURA 2 — CURVAS PR (todas as combinações, por classe)
# =========================================================
fig, axes = plt.subplots(1, 2, figsize=(14, 6))

for ci, cn in enumerate(class_names):
    ax = axes[ci]
    for mi, res in enumerate(resultados):
        p    = res['pr'][cn]
        ls   = line_styles[mi % len(line_styles)]
        ax.plot(p['rec'], p['prec'], lw=2, linestyle=ls,
                label=f'{res["nome"]} (AP={p["ap"]:.4f})')
        ax.scatter(p['rec_opt'], p['prec_opt'], zorder=5, s=60,
                   label=f'  Thr={p["thr_opt"]:.3f} P={p["prec_opt"]:.3f} R={p["rec_opt"]:.3f} F1={p["f1_opt"]:.3f}')
        ax.axhline(y=p['baseline'], lw=1, linestyle=':', alpha=0.4,
                   label=f'  Baseline={p["baseline"]:.3f}')

    ax.set_xlabel('Recall',    fontsize=11)
    ax.set_ylabel('Precision', fontsize=11)
    ax.set_title(f'Curva PR — Classe "{cn}"', fontsize=12)
    ax.legend(loc='lower left', fontsize=7)
    ax.grid(True, alpha=0.3)
    ax.set_xlim([0,1]); ax.set_ylim([0,1.05])

plt.tight_layout()
plt.savefig(os.path.join(save_path, 'pr_comparacao_late_fusion.png'), dpi=600, bbox_inches='tight')
plt.show()

# =========================================================
# IMPRESSÃO NO TERMINAL
# =========================================================
metricas_ordem = [
    'TP','FP','FN','TN',
    'Accuracy','Sensitivity','Specificity','Precision',
    'FalsePositiveRate','F1_score','MatthewsCorrelationCoef','Kappa',
]

for res in resultados:
    print('\n' + '='*65)
    print(f'MODELO: {res["nome"]}')
    print('='*65)
    for cn in class_names:
        print(f'\n  --- Classe positiva: "{cn}" ---')
        m = res['metricas'][cn]
        for k in metricas_ordem:
            v = m[k]
            print(f'    {k:<30}: {v}' if isinstance(v, int) else f'    {k:<30}: {v:.4f}')
        r = res['roc'][cn]
        p = res['pr'][cn]
        print(f'    {"AUC ROC":<30}: {r["auc"]:.4f}')
        print(f'    {"Threshold ROC (Youden)":<30}: {r["thr_opt"]:.4f}  Sens={r["sens_opt"]:.4f}  Espec={r["espec_opt"]:.4f}')
        print(f'    {"AP (PR-AUC)":<30}: {p["ap"]:.4f}')
        print(f'    {"Threshold PR (F1-max)":<30}: {p["thr_opt"]:.4f}  Prec={p["prec_opt"]:.4f}  Rec={p["rec_opt"]:.4f}  F1={p["f1_opt"]:.4f}')

# =========================================================
# DOCUMENTO WORD — TABELA COMPARATIVA
# =========================================================
doc = Document()

# Título
titulo = doc.add_heading('Comparação de Combinações — Late Fusion (Minimum Post-Hoc) — Dataset Gastrointestinal', level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Métricas calculadas no conjunto de teste para cada combinação de late fusion '
    '(média dos escores preditos das redes individuais), considerando cada classe '
    'como classe positiva (estratégia one-vs-rest): primeiro com "lesion" como '
    'positiva, depois com "normal" como positiva. '
    'Threshold ROC otimizado pelo índice de Youden. '
    'Threshold PR otimizado pelo F1-score máximo.'
)

# Colunas da tabela
col_metricas = metricas_ordem + [
    'AUC ROC', 'Threshold ROC', 'Sens@ROC', 'Espec@ROC',
    'AP (PR-AUC)', 'Threshold PR', 'Prec@PR', 'Rec@PR', 'F1@PR',
]

n_modelos = len(resultados)
n_classes  = len(class_names)

# Uma tabela por classe positiva (lesion primeiro, depois normal)
for cn in class_names:
    doc.add_heading(f'Classe positiva: {cn}', level=2)

    # cabeçalho: 1 coluna de métrica + 1 coluna por combinação
    table = doc.add_table(rows=1 + len(col_metricas), cols=1 + n_modelos)
    table.style = 'Table Grid'

    # Linha de cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = 'Métrica'
    hdr[0].paragraphs[0].runs[0].bold = True
    for mi, res in enumerate(resultados):
        hdr[mi+1].text = res['nome']
        hdr[mi+1].paragraphs[0].runs[0].bold = True
        hdr[mi+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linhas de dados
    for ri, met in enumerate(col_metricas):
        row = table.rows[ri+1].cells
        row[0].text = met
        row[0].paragraphs[0].runs[0].bold = True

        for mi, res in enumerate(resultados):
            m = res['metricas'][cn]
            r = res['roc'][cn]
            p = res['pr'][cn]

            if met in metricas_ordem:
                v = m[met]
                txt = str(v) if isinstance(v, int) else f'{v:.4f}'
            elif met == 'AUC ROC':
                txt = f'{r["auc"]:.4f}'
            elif met == 'Threshold ROC':
                txt = f'{r["thr_opt"]:.4f}'
            elif met == 'Sens@ROC':
                txt = f'{r["sens_opt"]:.4f}'
            elif met == 'Espec@ROC':
                txt = f'{r["espec_opt"]:.4f}'
            elif met == 'AP (PR-AUC)':
                txt = f'{p["ap"]:.4f}'
            elif met == 'Threshold PR':
                txt = f'{p["thr_opt"]:.4f}'
            elif met == 'Prec@PR':
                txt = f'{p["prec_opt"]:.4f}'
            elif met == 'Rec@PR':
                txt = f'{p["rec_opt"]:.4f}'
            elif met == 'F1@PR':
                txt = f'{p["f1_opt"]:.4f}'
            else:
                txt = ''

            row[mi+1].text = txt
            row[mi+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()   # espaço entre tabelas

# Inserir figuras ROC e PR
doc.add_heading('Curvas ROC', level=2)
doc.add_picture(os.path.join(save_path, 'roc_comparacao_late_fusion.png'), width=Inches(6.0))
doc.add_paragraph()

doc.add_heading('Curvas PR', level=2)
doc.add_picture(os.path.join(save_path, 'pr_comparacao_late_fusion.png'), width=Inches(6.0))

# Salvar
docx_path = os.path.join(save_path, 'comparacao_late_fusion.docx')
doc.save(docx_path)

print(f'\nDocumento Word salvo em : {docx_path}')
print(f'Curva ROC salva em      : {os.path.join(save_path, "roc_comparacao_late_fusion.png")}')
print(f'Curva PR  salva em      : {os.path.join(save_path, "pr_comparacao_late_fusion.png")}')